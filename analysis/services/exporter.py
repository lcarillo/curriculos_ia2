from django.http import HttpResponse
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_pdf(resume_text: str, filename: str) -> HttpResponse:
    """Exporta o currículo para PDF"""
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1
    )

    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=12
    )

    normal_style = styles['Normal']

    story = []
    lines = resume_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.isupper() and len(line) < 50:
            story.append(Paragraph(line, title_style))
        elif line.endswith(':'):
            story.append(Paragraph(line[:-1], heading_style))
        else:
            story.append(Paragraph(line, normal_style))

        story.append(Spacer(1, 12))

    doc.build(story)
    pdf = buffer.getvalue()
    buffer.close()
    response.write(pdf)

    return response


def export_docx(resume_text: str, filename: str) -> HttpResponse:
    """Exporta o currículo para DOCX"""
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'

    doc = docx.Document()

    lines = resume_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.isupper() and len(line) < 50:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.endswith(':'):
            p = doc.add_paragraph()
            run = p.add_run(line[:-1])
            run.bold = True
            run.font.size = Pt(14)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response.write(buffer.getvalue())
    buffer.close()

    return response

